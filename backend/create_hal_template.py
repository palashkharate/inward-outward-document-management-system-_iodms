"""
FR-143: Generate a professional HAL AURDC letterhead template (.docx).

This script creates a proper Word document template with:
- HAL logo in the header
- Company name and division details
- Reference number, date, To, Subject placeholders
- Body content area
- CC and signature block
- Footer with office address

All {{placeholder}} tags are replaced by the backend when a draft is created.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


def create_hal_template(output_path: str, logo_path: str = None):
    """FR-143: Creates a professional HAL letterhead .docx template."""
    
    doc = Document()
    
    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Cm(21)      # A4 width
    section.page_height = Cm(29.7)   # A4 height
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # --- Header with Logo ---
    header = section.header
    header.is_linked_to_previous = False
    
    # Add logo if available
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if logo_path and os.path.exists(logo_path):
        run = header_para.add_run()
        run.add_picture(logo_path, width=Inches(0.8))
        header_para.add_run("\n")
    
    # Company name
    run = header_para.add_run("HINDUSTAN AERONAUTICS LIMITED")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)  # Navy blue
    run.font.name = "Arial"
    
    header_para.add_run("\n")
    
    # Division name
    run = header_para.add_run("Aircraft Research & Design Centre, Nashik Division")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.font.name = "Arial"
    
    header_para.add_run("\n")
    
    # Separator line
    run = header_para.add_run("━" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # --- Reference & Date Line ---
    ref_para = doc.add_paragraph()
    ref_para.space_before = Pt(6)
    ref_para.space_after = Pt(2)
    run = ref_para.add_run("Ref: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run = ref_para.add_run("{{outward_reference}}")
    run.font.size = Pt(10)
    run.font.name = "Arial"
    
    date_para = doc.add_paragraph()
    date_para.space_before = Pt(0)
    date_para.space_after = Pt(12)
    run = date_para.add_run("Date: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run = date_para.add_run("{{date}}")
    run.font.size = Pt(10)
    run.font.name = "Arial"
    
    # --- To Block ---
    to_label = doc.add_paragraph()
    to_label.space_before = Pt(0)
    to_label.space_after = Pt(2)
    run = to_label.add_run("To,")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"
    
    to_addr = doc.add_paragraph()
    to_addr.space_before = Pt(0)
    to_addr.space_after = Pt(12)
    run = to_addr.add_run("{{to}}")
    run.font.size = Pt(11)
    run.font.name = "Arial"
    
    # --- Subject Line ---
    subj_para = doc.add_paragraph()
    subj_para.space_before = Pt(0)
    subj_para.space_after = Pt(12)
    run = subj_para.add_run("Sub: ")
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)
    run.font.name = "Arial"
    run = subj_para.add_run("{{subject}}")
    run.underline = True
    run.font.size = Pt(11)
    run.font.name = "Arial"
    
    # --- Salutation ---
    sal_para = doc.add_paragraph()
    sal_para.space_before = Pt(0)
    sal_para.space_after = Pt(12)
    run = sal_para.add_run("Sir / Madam,")
    run.font.size = Pt(11)
    run.font.name = "Arial"
    
    # --- Body Placeholder ---
    body_para = doc.add_paragraph()
    body_para.space_before = Pt(0)
    body_para.space_after = Pt(24)
    run = body_para.add_run("{{body}}")
    run.font.size = Pt(11)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(150, 150, 150)
    
    # --- CC Line ---
    cc_para = doc.add_paragraph()
    cc_para.space_before = Pt(12)
    cc_para.space_after = Pt(6)
    run = cc_para.add_run("CC: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run = cc_para.add_run("{{cc}}")
    run.font.size = Pt(10)
    run.font.name = "Arial"
    
    # --- Signature Block (right-aligned) ---
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_para.space_before = Pt(36)
    sig_para.space_after = Pt(2)
    run = sig_para.add_run("Yours faithfully,")
    run.font.size = Pt(11)
    run.font.name = "Arial"
    
    sig_name = doc.add_paragraph()
    sig_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_name.space_before = Pt(24)
    sig_name.space_after = Pt(0)
    run = sig_name.add_run("{{prepared_by}}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"
    
    sig_title = doc.add_paragraph()
    sig_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_title.space_before = Pt(0)
    sig_title.space_after = Pt(0)
    run = sig_title.add_run("For General Manager")
    run.font.size = Pt(10)
    run.font.name = "Arial"
    
    sig_org = doc.add_paragraph()
    sig_org.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_org.space_before = Pt(0)
    run = sig_org.add_run("HAL, AURDC Nashik")
    run.font.size = Pt(10)
    run.font.name = "Arial"
    
    # --- Remarks (small, at bottom) ---
    rem_para = doc.add_paragraph()
    rem_para.space_before = Pt(12)
    run = rem_para.add_run("Remarks: ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(100, 100, 100)
    run = rem_para.add_run("{{remarks}}")
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # --- Footer ---
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("━" * 60)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0, 51, 102)
    footer_para.add_run("\n")
    run = footer_para.add_run(
        "HAL AURDC, Ojhar Township, Nashik — 422 207, Maharashtra, India  |  "
        "Phone: 0253-2384000  |  www.hal-india.co.in"
    )
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.font.name = "Arial"
    
    # --- Save ---
    dirname = os.path.dirname(output_path)
    if dirname:
        os.makedirs(dirname, exist_ok=True)
    doc.save(output_path)
    print(f"Template created: {output_path}")
    return output_path


if __name__ == "__main__":
    # FR-143: Quick test — generate template in current directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    logo = os.path.join(script_dir, "hal_logo.jpg")
    create_hal_template("hal_test_template.docx", logo)
    print("Done! Open hal_test_template.docx in Word to verify.")
