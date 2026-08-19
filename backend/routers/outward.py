import datetime
import os
import shutil
import json
import uuid
from urllib.parse import quote
from .link_utils import sync_bidirectional_links
from typing import Optional, List
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, Request
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sqlalchemy.exc import IntegrityError
from pydantic import BaseModel

import models
from database import get_db, get_iodms_root_path, get_iodms_lan_share_path
from routers.inward import get_effective_year
import filesystem_utils
from auth_utils import get_current_user

router = APIRouter(dependencies=[Depends(get_current_user)])

# --- Pydantic Request Models ---
class DraftCreate(BaseModel):
    outward_no: Optional[str] = None
    folder_id: str
    issuing_date: str
    address_to: List[int]
    cc_to: List[int]
    subject: str
    remarks: Optional[str] = ""
    prepared_by: str
    actioned_by: str
    template_type: str
    target_year: Optional[int] = None
    linked_documents: Optional[List[str]] = []
    document_body: Optional[dict] = None

class DraftLockAction(BaseModel):
    user_id: Optional[str] = None

# --- Helper functions ---

SUPPORTED_TEMPLATE_EXTENSIONS = {".doc", ".docx", ".rtf", ".txt"}


def get_selected_template(template_id, db: Session):
    """Returns the selected template when the submitted value is a template ID."""
    try:
        return db.query(models.DocumentTemplate).filter(
            models.DocumentTemplate.id == int(template_id)
        ).first()
    except (TypeError, ValueError):
        return None


def get_template_document_extension(template_id, db: Session) -> str:
    """Uses the template's actual format so copied templates stay valid files."""
    template = get_selected_template(template_id, db)
    if not template:
        return ".docx"
    extension = os.path.splitext(template.file_path)[1].lower()
    return extension if extension in SUPPORTED_TEMPLATE_EXTENSIONS else ".docx"

def check_draft_locks(db: Session):
    """Auto-expires locks older than 30 minutes."""
    thirty_mins_ago = datetime.datetime.now() - datetime.timedelta(minutes=30)
    db.query(models.DraftFile).filter(
        models.DraftFile.is_locked == True,
        models.DraftFile.locked_at < thirty_mins_ago
    ).update({"is_locked": False, "locked_by": None, "locked_at": None})
    db.commit()

# FR-052: Build the direct LAN path officers open in Microsoft Word.
def build_lan_document_open_info(relative_path: str, draft_id: int, request: Request) -> dict:
    lan_root = get_iodms_lan_share_path()
    if not lan_root or not relative_path:
        return {
            "lan_shared_path": None,
            "lan_file_uri": None,
            "word_launcher_uri": None,
            "word_open_uri": None
        }

    normalized_relative = str(relative_path).strip().replace("/", "\\").lstrip("\\")
    lan_shared_path = lan_root.rstrip("\\/") + "\\" + normalized_relative
    file_url_path = quote(lan_shared_path.replace("\\", "/").lstrip("/"), safe="/:")
    lan_file_uri = (
        f"file://{file_url_path}"
        if lan_shared_path.startswith("\\\\")
        else f"file:///{file_url_path}"
    )
    return {
        "lan_shared_path": lan_shared_path,
        "lan_file_uri": lan_file_uri,
        "word_launcher_uri": f"iodms-word://open?path={quote(lan_shared_path, safe='')}",
        "word_open_uri": f"ms-word:ofe|u|{request.url.scheme}://{request.headers.get('host', 'localhost')}/api/webdav/drafts/{draft_id}/{os.path.basename(lan_shared_path.replace('\\\\\\\\', '/').replace('\\', '/'))}"
    }

def log_edit(db: Session, record_type: str, record_id: str, action: str, user_id: str, changes: dict = None):
    """Helper to add an entry to the EditLog."""
    log = models.EditLog(
        record_type=record_type,
        record_id=str(record_id),
        action=action,
        changes=changes,
        edited_by=user_id
    )
    db.add(log)
    db.commit()

def parse_document_date(value: str) -> datetime.date:
    """Parses and bounds office document dates before saving to registers."""
    try:
        parsed = datetime.datetime.strptime(value, "%Y-%m-%d").date()
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format. Use YYYY-MM-DD")

    min_date = datetime.date(1947, 8, 15)
    today = datetime.date.today()
    if parsed < min_date or parsed > today:
        raise HTTPException(
            status_code=400,
            detail=f"Date of Document must be between {min_date.isoformat()} and {today.isoformat()}."
        )
    return parsed

def ensure_draft_not_pending_deletion(db: Session, draft_id: int):
    pending = db.query(models.PendingDeletion).filter(
        models.PendingDeletion.source_table == "draft_files",
        models.PendingDeletion.record_id == str(draft_id),
        models.PendingDeletion.status == "Pending"
    ).first()
    if pending:
        raise HTTPException(status_code=400, detail="Deletion is pending for this draft. Admin approval is required before further action.")

# FR-055: Pre-assign and reserve Outward number
def get_next_outward_no(folder_id: str, year: int, db: Session) -> str:
    """Gets the next sequential Outward Number by looking at the whole outward year.
    
    Implements:
    - FR-055: Assigns outward numbers only when a draft is dispatched.
      Outward numbers are yearly global numbers; Folder ID only controls storage grouping.
    """
    # 1. Fetch from outward_register for the whole year, not per Folder ID.
    register_nos = db.query(models.OutwardRegister.outward_no).filter(
        models.OutwardRegister.year == year
    ).all()
    
    numbers = []
    for (no_str,) in register_nos:
        try:
            numbers.append(int(no_str))
        except ValueError:
            pass
            
    if not numbers:
        return "001"
        
    next_val = max(numbers) + 1
    if next_val <= 999:
        return f"{next_val:03d}"
    else:
        return str(next_val)


def build_outward_reference(data: dict) -> str:
    outward_no = data.get("outward_no")
    folder_id = data.get("folder_id") or ""
    year = data.get("year") or ""
    if not outward_no or str(outward_no).upper().startswith(("DRAFT", "PENDING")):
        return f"HAL/NK/D/DAE/{folder_id}/{year}/Pending Dispatch"
    return f"HAL/NK/D/DAE/{folder_id}/{year}/{outward_no}"


def get_address_text(db: Session, address_ids: list[int]) -> str:
    if not address_ids:
        return ""
    addr = db.query(models.AddressBook).filter(models.AddressBook.address_id == address_ids[0]).first()
    if not addr:
        return ""
    return "\n".join(filter(None, [
        addr.name,
        addr.designation,
        addr.organisation,
        addr.address_line_1,
        addr.address_line_2
    ]))


def get_cc_text(db: Session, cc_ids: list[int]) -> str:
    names = []
    for cc_id in cc_ids or []:
        addr = db.query(models.AddressBook).filter(models.AddressBook.address_id == cc_id).first()
        if addr:
            names.append(addr.name)
    return ", ".join(names)


def build_template_replacements(data: dict, db: Session) -> dict:
    return {
        "outward_reference": build_outward_reference(data),
        "folder_id": data.get("folder_id") or "",
        "year": data.get("year") or "",
        "outward_no": data.get("outward_no") or "",
        "subject": data.get("subject") or "",
        "date": data.get("issuing_date") or datetime.date.today().isoformat(),
        "prepared_by": data.get("prepared_by") or "",
        "to": get_address_text(db, data.get("address_to") or []),
        "cc": get_cc_text(db, data.get("cc_to") or []),
        "remarks": data.get("remarks") or ""
    }


def replace_docx_placeholders(filepath: str, replacements: dict):
    from docx import Document

    doc = Document(filepath)
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for paragraph in paragraphs:
        for key, value in replacements.items():
            token = "{{" + key + "}}"
            if token in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(token, str(value or ""))
    doc.save(filepath)


def replace_text_placeholders(filepath: str, replacements: dict):
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        for key, value in replacements.items():
            content = content.replace("{{" + key + "}}", str(value or ""))
        if replacements.get("outward_reference"):
            content = content.replace(
                "Draft - outward number will be assigned on dispatch",
                str(replacements["outward_reference"])
            )
        if replacements.get("outward_no"):
            content = content.replace("Pending Dispatch", str(replacements["outward_no"]))
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
    except Exception:
        pass


def is_rtf_or_text_word_file(filepath: str) -> bool:
    try:
        with open(filepath, "rb") as f:
            start = f.read(16).lstrip()
        return start.startswith(b"{\\rtf") or not start.startswith(b"\xd0\xcf\x11\xe0")
    except Exception:
        return False


def stamp_outward_reference(filepath: str, data: dict, db: Session):
    if not os.path.exists(filepath):
        return
    try:
        lower = filepath.lower()
        if lower.endswith(".docx"):
            replace_docx_placeholders(filepath, build_template_replacements(data, db))
        elif lower.endswith((".doc", ".rtf", ".txt")) and is_rtf_or_text_word_file(filepath):
            replace_text_placeholders(filepath, build_template_replacements(data, db))
    except Exception:
        pass





# FR-042: Generate Word document draft with placeholder tags
def create_draft_document(filepath: str, data: dict, db: Session):
    """Creates a document from template or fallback basic formatted text file.
    
    Implements:
    - FR-143: Generates a document using uploaded DocumentTemplates.
    """
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    template = get_selected_template(data.get("template_type"), db)

    document_body = data.get("document_body")

    if template:
        src_path = os.path.join(get_iodms_root_path(), template.file_path)
        if os.path.exists(src_path):
            try:
                src_ext = os.path.splitext(src_path)[1].lower()
                if src_ext == ".docx" and filepath.lower().endswith(".docx"):
                    shutil.copyfile(src_path, filepath)
                    replace_docx_placeholders(filepath, build_template_replacements(data, db))
                    
                    if document_body and "blocks" in document_body:
                        import docx
                        doc = docx.Document(filepath)
                        # Append editor blocks to the doc
                        for block in document_body["blocks"]:
                            if block["type"] == "paragraph":
                                doc.add_paragraph(block["data"].get("text", ""))
                            elif block["type"] == "header":
                                doc.add_heading(block["data"].get("text", ""), level=block["data"].get("level", 1))
                            elif block["type"] == "list":
                                style = 'List Number' if block["data"].get("style") == "ordered" else 'List Bullet'
                                for item in block["data"].get("items", []):
                                    doc.add_paragraph(item, style=style)
                        doc.save(filepath)
                    return
                if src_ext in [".doc", ".rtf", ".txt"]:
                    shutil.copyfile(src_path, filepath)
                    if is_rtf_or_text_word_file(filepath):
                        replace_text_placeholders(filepath, build_template_replacements(data, db))
                    # Fallback text appending for non-docx
                    if document_body and "blocks" in document_body:
                        with open(filepath, "a", encoding="utf-8") as f:
                            f.write("\n\n")
                            for block in document_body["blocks"]:
                                if block["type"] in ["paragraph", "header"]:
                                    f.write(block["data"].get("text", "") + "\n\n")
                                elif block["type"] == "list":
                                    for item in block["data"].get("items", []):
                                        f.write(f"- {item}\n")
                                    f.write("\n")
                    return
            except Exception as e:
                print(f"Template parsing failed: {e}")
                pass
            
    # If no template or copying failed, generate basic file
    import docx
    doc = docx.Document()
    doc.add_heading(data.get("subject", "Document"), 0)
    
    if document_body and "blocks" in document_body:
        for block in document_body["blocks"]:
            if block["type"] == "paragraph":
                doc.add_paragraph(block["data"].get("text", ""))
            elif block["type"] == "header":
                doc.add_heading(block["data"].get("text", ""), level=block["data"].get("level", 1))
            elif block["type"] == "list":
                style = 'List Number' if block["data"].get("style") == "ordered" else 'List Bullet'
                for item in block["data"].get("items", []):
                    doc.add_paragraph(item, style=style)
    else:
        doc.add_paragraph("[Place your letter body contents here...]")
        doc.add_paragraph("Remarks: " + str(data.get("remarks") or ""))
        
    doc.save(filepath)


def is_blank_fallback_draft(filepath: str) -> bool:
    """Recognizes the old generic file produced when a DOCX template was skipped."""
    try:
        with open(filepath, "rb") as file_handle:
            contents = file_handle.read()
        return b"[Place your letter body contents here...]" in contents
    except OSError:
        return False


def rebuild_blank_draft_from_template(draft, db: Session) -> bool:
    """Repairs only known blank legacy fallbacks; user-edited files are never replaced."""
    template_extension = get_template_document_extension(draft.template_type, db)
    current_extension = os.path.splitext(draft.file_path)[1].lower()
    if current_extension == template_extension:
        return False

    old_full_path = os.path.join(get_iodms_root_path(), draft.file_path)
    if not is_blank_fallback_draft(old_full_path):
        return False

    new_relative_path = os.path.splitext(draft.file_path)[0] + template_extension
    new_full_path = os.path.join(get_iodms_root_path(), new_relative_path)
    draft_data = {
        "outward_no": draft.outward_no,
        "folder_id": draft.folder_id,
        "issuing_date": draft.issuing_date.isoformat(),
        "address_to": draft.address_to or [],
        "cc_to": draft.cc_to or [],
        "subject": draft.subject,
        "remarks": draft.remarks,
        "prepared_by": draft.prepared_by,
        "actioned_by": draft.actioned_by,
        "template_type": draft.template_type,
        "linked_documents": draft.linked_documents or [],
        "document_body": draft.document_body,
        "year": draft.year,
    }
    create_draft_document(new_full_path, draft_data, db)
    draft.file_path = new_relative_path.replace("\\", "/")
    draft.attachment_paths = [
        draft.file_path if path == os.path.splitext(new_relative_path)[0] + current_extension else path
        for path in (draft.attachment_paths or [])
    ]
    if os.path.exists(old_full_path):
        os.remove(old_full_path)
    return True


# --- Endpoints ---

# FR-055: Preview next Outward No. without reserving it
@router.get("/next-no")
def get_next_no(
    folder_id: str, 
    target_year: Optional[int] = None, 
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Previews the next available Outward number without reserving it.
    
    Implements:
    - FR-055: Outward numbers are assigned only at dispatch.
    - FR-144: Support target_year override for previous year entries.
    """
    year = target_year if target_year else get_effective_year()
    return {"outward_no": get_next_outward_no(folder_id, year, db), "year": year, "reserved": False}


# FR-042: Save Draft
@router.post("/draft")
def save_draft(
    payload: DraftCreate,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Saves outward details as a draft using the selected template's file format.
    
    Implements:
    - FR-042: Generates a draft file under IODMS/Drafts/{Year}/{FolderID}/draft-...<template extension>
    - FR-144: Support target_year override
    """
    year = payload.target_year if payload.target_year else get_effective_year()
    actor_id = current_user.get("user_id")

    # Preserve the selected template's extension; copying a DOCX to a .doc
    # filename caused Word to receive a generic fallback instead of the template.
    timestamp = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
    draft_marker = "DRAFT"
    document_extension = get_template_document_extension(payload.template_type, db)
    filename = f"draft-{actor_id}-{timestamp}{document_extension}"
    
    relative_folder, full_folder = filesystem_utils.ensure_folder_path(get_iodms_root_path(), "Drafts", year, payload.folder_id)
    relative_path = os.path.join(relative_folder, filename).replace("\\", "/")
    full_path = os.path.join(full_folder, filename)
    
    payload_dict = payload.model_dump()
    payload_dict["outward_no"] = None
    payload_dict["year"] = year
    payload_dict["actioned_by"] = actor_id
    
    # Save the physical file on disk (FR-042)
    try:
        create_draft_document(full_path, payload_dict, db)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create draft document on disk: {str(e)}")

    iss_date = parse_document_date(payload.issuing_date)

    draft_record = models.DraftFile(
        file_path=relative_path,
        outward_no=draft_marker,
        folder_id=payload.folder_id,
        issuing_date=iss_date,
        address_to=payload.address_to,
        cc_to=payload.cc_to,
        subject=payload.subject,
        remarks=payload.remarks,
        prepared_by=payload.prepared_by,
        actioned_by=actor_id,
        template_type=payload.template_type,
        linked_documents=payload.linked_documents,
        attachment_paths=[relative_path],
        document_body=payload.document_body,
        year=year
    )
    db.add(draft_record)
    db.flush()
    log_edit(db, "draft", str(draft_record.draft_id), "create", actor_id, payload_dict)

    db.commit()
    return {"message": "Draft created successfully. Outward number will be assigned on dispatch.", "draft_id": draft_record.draft_id, "outward_no": None, "success": True}


# FR-044, FR-052: Update Draft
@router.put("/drafts/{draft_id}")
def update_draft(
    draft_id: int,
    payload: DraftCreate,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Updates an existing draft document."""
    draft = db.query(models.DraftFile).filter(models.DraftFile.draft_id == draft_id).first()
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    ensure_draft_not_pending_deletion(db, draft_id)
    
    actor_id = current_user.get("user_id")

    iss_date = parse_document_date(payload.issuing_date)

    draft.issuing_date = iss_date
    draft.folder_id = payload.folder_id
    draft.address_to = payload.address_to
    draft.cc_to = payload.cc_to
    draft.subject = payload.subject
    draft.remarks = payload.remarks
    draft.prepared_by = payload.prepared_by
    draft.actioned_by = actor_id
    draft.template_type = payload.template_type
    draft.linked_documents = payload.linked_documents
    if payload.document_body is not None:
        draft.document_body = payload.document_body

    # Auto-upgrade legacy .doc drafts to .docx to prevent python-docx zip corruption
    if draft.file_path.lower().endswith(".doc"):
        old_full_path = os.path.join(get_iodms_root_path(), draft.file_path)
        draft.file_path = draft.file_path[:-4] + ".docx"
        if os.path.exists(old_full_path):
            try:
                os.remove(old_full_path)
            except OSError:
                pass
        
        # Update attachment references if the primary document was renamed
        if draft.attachment_paths:
            new_paths = []
            for p in draft.attachment_paths:
                if p.lower().endswith(".doc") and p[:-4] == old_full_path[:-4]:
                    new_paths.append(draft.file_path)
                else:
                    new_paths.append(p)
            draft.attachment_paths = new_paths

    full_path = os.path.join(get_iodms_root_path(), draft.file_path)
    
    try:
        payload_dict = payload.model_dump()
        payload_dict["outward_no"] = None
        payload_dict["folder_id"] = payload.folder_id
        payload_dict["year"] = draft.year
        create_draft_document(full_path, payload_dict, db)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to overwrite file on disk: {str(e)}")

    changes = payload.model_dump()
    changes["actioned_by"] = actor_id
    log_edit(db, "draft", str(draft_id), "edit", actor_id, changes)

    db.commit()
    return {"message": "Draft updated successfully", "success": True}


# FR-170b: Attach supporting files to an outward draft from Compose Outward
@router.post("/drafts/{draft_id}/attachments")
def attach_draft_files(
    draft_id: int,
    files: List[UploadFile] = File([]),
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Adds supporting PDFs, PPTs, DOCs, or other office files to an existing draft."""
    draft = db.query(models.DraftFile).filter(models.DraftFile.draft_id == draft_id).first()
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    ensure_draft_not_pending_deletion(db, draft_id)
    if not files or (len(files) == 1 and files[0].filename == ""):
        raise HTTPException(status_code=400, detail="At least one file is required.")

    actor_id = current_user.get("user_id")
    relative_folder, full_folder = filesystem_utils.ensure_folder_path(
        get_iodms_root_path(), "Drafts", draft.year, draft.folder_id
    )
    existing_paths = draft.attachment_paths or []
    new_paths = []

    for idx, file in enumerate(files, start=1):
        if not file.filename:
            continue
        ext = os.path.splitext(file.filename)[1] or ".bin"
        filename = f"{draft.outward_no}_attachment_{len(existing_paths) + idx}{ext}"
        relative_path = os.path.join(relative_folder, filename).replace("\\", "/")
        full_path = os.path.join(full_folder, filename)

        try:
            with open(full_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            final_abs_path, was_compressed = filesystem_utils.compress_file_if_large(full_path)
            if was_compressed:
                final_filename = os.path.basename(final_abs_path)
                relative_path = os.path.join(relative_folder, final_filename).replace("\\", "/")
            new_paths.append(relative_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to save attachment: {str(e)}")

    draft.attachment_paths = existing_paths + new_paths
    log_edit(db, "draft", str(draft_id), "attach", actor_id, {"files": new_paths})
    db.commit()
    return {"message": "Draft attachment files uploaded successfully", "files": new_paths, "success": True}


# FR-170c: Remove a supporting file from an active outward draft
@router.delete("/drafts/{draft_id}/attachments")
def delete_draft_attachment(
    draft_id: int,
    path: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Deletes a supporting attachment while protecting the main draft document."""
    draft = db.query(models.DraftFile).filter(models.DraftFile.draft_id == draft_id).first()
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    ensure_draft_not_pending_deletion(db, draft_id)
    if path == draft.file_path:
        raise HTTPException(status_code=400, detail="Main draft document cannot be deleted. Replace it through Open / Edit and re-upload.")

    existing_paths = draft.attachment_paths or []
    if path not in existing_paths:
        raise HTTPException(status_code=404, detail="Attachment not found on this draft.")

    actor_id = current_user.get("user_id")
    root_path = os.path.abspath(get_iodms_root_path())
    full_path = os.path.abspath(os.path.join(root_path, path))
    if os.path.commonpath([root_path, full_path]) != root_path:
        raise HTTPException(status_code=400, detail="Invalid attachment path.")

    draft.attachment_paths = [p for p in existing_paths if p != path]
    if os.path.exists(full_path):
        try:
            os.remove(full_path)
        except OSError as e:
            raise HTTPException(status_code=500, detail=f"Failed to delete attachment file: {str(e)}")

    log_edit(db, "draft", str(draft_id), "delete_attachment", actor_id, {"file": path})
    db.commit()
    return {"message": "Attachment deleted successfully.", "success": True}


# FR-057: Direct Draft Upload
@router.post("/drafts/upload")
def upload_existing_draft(
    folder_id: str = Form(...),
    issuing_date: str = Form(...),
    address_to: str = Form(...), # comma separated IDs
    cc_to: str = Form(""),       # comma separated IDs
    subject: str = Form(...),
    remarks: str = Form(""),
    prepared_by: str = Form(...),
    actioned_by: str = Form(...),
    files: List[UploadFile] = File([]),
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Uploads one or more existing files directly as a new draft.
    
    Implements:
    - FR-057: Bypasses template generation and uses user-uploaded file.
    """
    year = get_effective_year()
    actor_id = current_user.get("user_id")
    timestamp = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
    draft_marker = "DRAFT"
    
    attachment_paths = []
    any_compressed = False

    if not files or (len(files) == 1 and files[0].filename == ""):
        raise HTTPException(status_code=400, detail="At least one file is required.")

    for idx, file in enumerate(files):
        if not file.filename: continue
        ext = os.path.splitext(file.filename)[1]
        
        if len(files) > 1:
            filename = f"draft-{actor_id}-{timestamp}_{idx+1}{ext}"
        else:
            filename = f"draft-{actor_id}-{timestamp}{ext}"
        
        relative_folder, full_folder = filesystem_utils.ensure_folder_path(get_iodms_root_path(), "Drafts", year, folder_id)
        relative_path = os.path.join(relative_folder, filename).replace("\\", "/")
        full_path = os.path.join(full_folder, filename)
        
        try:
            with open(full_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
                
            final_abs_path, was_compressed = filesystem_utils.compress_file_if_large(full_path)
            if was_compressed: any_compressed = True
            
            if was_compressed:
                final_filename = os.path.basename(final_abs_path)
                attachment_paths.append(os.path.join(relative_folder, final_filename).replace("\\", "/"))
            else:
                attachment_paths.append(relative_path)
                
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to save uploaded file: {str(e)}")

    iss_date = parse_document_date(issuing_date)
        
    addr_list = [int(i.strip()) for i in address_to.split(",") if i.strip().isdigit()]
    cc_list = [int(i.strip()) for i in cc_to.split(",") if i.strip().isdigit()]

    new_draft = models.DraftFile(
        file_path=attachment_paths[0] if attachment_paths else "",
        attachment_paths=attachment_paths,
        outward_no=draft_marker,
        folder_id=folder_id,
        issuing_date=iss_date,
        address_to=addr_list,
        cc_to=cc_list,
        subject=subject,
        remarks=remarks,
        prepared_by=prepared_by,
        actioned_by=actor_id,
        template_type="Manual Upload",
        is_locked=False,
        year=year
    )
    
    db.add(new_draft)
    db.commit()
    log_edit(db, "draft", str(new_draft.draft_id), "create", actor_id, {
        "outward_no": None,
        "folder_id": folder_id,
        "subject": subject,
        "uploaded_files": attachment_paths
    })
    return {"message": "Draft uploaded successfully. Outward number will be assigned on dispatch.", "draft_id": new_draft.draft_id, "outward_no": None, "success": True}

# FR-052: Re-Upload Draft File (after editing)
@router.put("/drafts/{draft_id}/reupload")
def reupload_draft_file(
    draft_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Re-uploads an edited draft file and releases its lock.
    
    Implements:
    - FR-052: User edits file locally and re-uploads it here.
    """
    draft = db.query(models.DraftFile).filter(models.DraftFile.draft_id == draft_id).first()
    actor_id = current_user.get("user_id")
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    ensure_draft_not_pending_deletion(db, draft_id)
        
    if not draft.is_locked:
        raise HTTPException(status_code=400, detail="Draft is not locked. Lock it first before re-uploading.")
    if draft.locked_by and draft.locked_by != actor_id:
        raise HTTPException(status_code=403, detail=f"This draft is locked by {draft.locked_by}. Ask them or an Admin to release the lock.")

    full_path = os.path.join(get_iodms_root_path(), draft.file_path)
    
    try:
        with open(full_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        # FR-161: Compress file if larger than 50MB
        final_abs_path, was_compressed = filesystem_utils.compress_file_if_large(full_path)
        if was_compressed:
            final_filename = os.path.basename(final_abs_path)
            # Update path in DB to point to new .zip file
            draft.file_path = os.path.join(os.path.dirname(draft.file_path), final_filename).replace("\\", "/")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to overwrite draft file: {str(e)}")

    # Release lock after successful upload
    draft.is_locked = False
    draft.locked_by = None
    draft.locked_at = None
    
    log_edit(db, "draft", str(draft_id), "reupload", actor_id, {"filename": file.filename})
    
    db.commit()
    
    return {"message": "Draft updated and unlocked successfully", "success": True}


# FR-044, FR-045: Modify Outward (Modify Mode)
@router.put("/modify/{folder_id}/{year}/{outward_no}")
def modify_outward(
    folder_id: str,
    year: int,
    outward_no: str,
    payload: DraftCreate,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Modifies an already dispatched outward record in place and updates its file.
    
    Implements:
    - FR-044: Pre-filled form edit. Modifies existing outward_register record and replaces file on disk.
    """
    record = db.query(models.OutwardRegister).filter(
        models.OutwardRegister.folder_id == folder_id,
        models.OutwardRegister.year == year,
        models.OutwardRegister.outward_no == outward_no
    ).first()

    if not record:
        raise HTTPException(status_code=404, detail="Outward record not found")
    actor_id = current_user.get("user_id")

    iss_date = parse_document_date(payload.issuing_date)

    # Update database record
    record.issuing_date = iss_date
    record.address_to = payload.address_to
    record.cc_to = payload.cc_to
    record.subject = payload.subject
    record.remarks = payload.remarks
    record.prepared_by = payload.prepared_by
    record.actioned_by = actor_id
    record.template_type = payload.template_type
    old_links = record.linked_documents or []
    record.linked_documents = payload.linked_documents
    if payload.document_body is not None:
        record.document_body = payload.document_body
    
    source_id = f"outward:{folder_id}:{year}:{outward_no}"
    sync_bidirectional_links(db, source_id, old_links, payload.linked_documents)

    # Recreate the file on disk
    full_path = os.path.join(get_iodms_root_path(), record.document_path)
    try:
        payload_dict = payload.model_dump()
        payload_dict["outward_no"] = outward_no
        payload_dict["folder_id"] = folder_id
        payload_dict["year"] = year
        create_draft_document(full_path, payload_dict, db)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to overwrite file on disk: {str(e)}")

    changes = payload.model_dump()
    changes["actioned_by"] = actor_id
    log_edit(db, "outward", f"{folder_id}:{year}:{outward_no}", "edit", actor_id, changes)

    db.commit()
    return {"message": "Outward record updated successfully", "success": True}


# FR-050: Get Drafts list
@router.get("/drafts")
def get_drafts(db: Session = Depends(get_db)):
    """Retrieves all drafts waiting to be dispatched.
    
    Implements:
    - FR-050: Lists all drafts.
    - FR-057: No per-user filtering (all users see all drafts).
    - Marks drafts pending deletion so users can see them greyed out until Admin action.
    """
    check_draft_locks(db)

    pending_deletes = db.query(models.PendingDeletion).filter(
        models.PendingDeletion.source_table == "draft_files",
        models.PendingDeletion.status == "Pending"
    ).all()
    pending_ids = {int(pd.record_id) for pd in pending_deletes if str(pd.record_id).isdigit()}

    # Hide legacy reservation placeholders from builds created before outward
    # numbers moved to dispatch time.
    query = db.query(models.DraftFile).filter(models.DraftFile.file_path != "[Reserved]")
    drafts = query.all()
    
    output = []
    for d in drafts:
        folder = db.query(models.FolderType).filter(models.FolderType.folder_id == d.folder_id).first()
        folder_name = folder.folder_name if folder else ""
        
        # Primary recipient name
        recipient_name = ""
        if d.address_to:
            addr = db.query(models.AddressBook).filter(models.AddressBook.address_id == d.address_to[0]).first()
            recipient_name = addr.name if addr else ""

        lan_open_info = build_lan_document_open_info(d.file_path, d.draft_id, request)
        output.append({
            "draft_id": d.draft_id,
            "file_path": d.file_path,
            "lan_shared_path": lan_open_info["lan_shared_path"],
            "lan_file_uri": lan_open_info["lan_file_uri"],
            "word_launcher_uri": lan_open_info["word_launcher_uri"],
            "word_open_uri": lan_open_info["word_open_uri"],
            "outward_no": d.outward_no,
            "folder_id": d.folder_id,
            "folder_name": folder_name,
            "issuing_date": d.issuing_date.isoformat(),
            "address_to": d.address_to,
            "recipient_name": recipient_name,
            "cc_to": d.cc_to,
            "subject": d.subject,
            "remarks": d.remarks,
            "prepared_by": d.prepared_by,
            "actioned_by": d.actioned_by,
            "attachment_paths": d.attachment_paths or [],
            "template_type": d.template_type,
            "is_locked": d.is_locked,
            "locked_by": d.locked_by,
            "created_on": d.created_on.isoformat(),
            "is_pending_deletion": d.draft_id in pending_ids
        })
    return output


# FR-052: Get specific draft
@router.get("/drafts/{draft_id}")
def get_draft(draft_id: int, request: Request, db: Session = Depends(get_db)):
    """Retrieves a specific draft for editing."""
    draft = db.query(models.DraftFile).filter(models.DraftFile.draft_id == draft_id).first()
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    
    # Format the payload for the frontend
    lan_open_info = build_lan_document_open_info(draft.file_path, draft.draft_id, request)
    return {
        "draft_id": draft.draft_id,
        "file_path": draft.file_path,
        "lan_shared_path": lan_open_info["lan_shared_path"],
        "lan_file_uri": lan_open_info["lan_file_uri"],
        "word_launcher_uri": lan_open_info["word_launcher_uri"],
        "word_open_uri": lan_open_info["word_open_uri"],
        "folder_id": draft.folder_id,
        "issuing_date": draft.issuing_date.isoformat(),
        "address_to": draft.address_to,
        "cc_to": draft.cc_to,
        "subject": draft.subject,
        "remarks": draft.remarks,
        "prepared_by": draft.prepared_by,
        "actioned_by": draft.actioned_by,
        "template_type": draft.template_type,
        "linked_documents": draft.linked_documents or [],
        "document_body": draft.document_body,
        "year": draft.year,
        "outward_no": draft.outward_no,
        "is_locked": draft.is_locked,
        "locked_by": draft.locked_by
    }
# FR-052: Lock draft file for editing
@router.put("/drafts/{draft_id}/lock")
def lock_draft(
    request: Request,
    draft_id: int,
    payload: Optional[DraftLockAction] = None,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Locks the draft to prevent editing conflicts.
    
    Implements:
    - FR-052: Checks if locked by another user and rejects request.
    """
    draft = db.query(models.DraftFile).filter(models.DraftFile.draft_id == draft_id).first()
    actor_id = current_user.get("user_id")
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    ensure_draft_not_pending_deletion(db, draft_id)

    if draft.is_locked and draft.locked_by != actor_id:
        # Find locked user name
        locker = db.query(models.User).filter(models.User.user_id == draft.locked_by).first()
        locker_name = locker.name if locker else draft.locked_by
        raise HTTPException(
            status_code=400,
            detail=f"This draft is currently being edited by {locker_name}. Please try again later."
        )

    repaired_from_template = rebuild_blank_draft_from_template(draft, db)

    draft.is_locked = True
    draft.locked_by = actor_id
    draft.locked_at = datetime.datetime.now()
    
    log_edit(db, "draft", str(draft_id), "lock", actor_id, {
        "repaired_from_template": repaired_from_template
    })
    
    db.commit()
    lan_open_info = build_lan_document_open_info(draft.file_path, draft.draft_id, request)
    return {
        "message": "Draft file locked for editing",
        "file_path": draft.file_path,
        "lan_shared_path": lan_open_info["lan_shared_path"],
        "lan_file_uri": lan_open_info["lan_file_uri"],
        "word_launcher_uri": lan_open_info["word_launcher_uri"],
        "word_open_uri": lan_open_info["word_open_uri"],
        "success": True
    }


# FR-053: Unlock draft file
@router.put("/drafts/{draft_id}/unlock")
def unlock_draft(
    draft_id: int,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Unlocks the draft file.
    
    Implements:
    - FR-053: Reset lock settings. Manual release available for Admin.
    """
    draft = db.query(models.DraftFile).filter(models.DraftFile.draft_id == draft_id).first()
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    ensure_draft_not_pending_deletion(db, draft_id)

    actor_id = current_user.get("user_id")
    if draft.locked_by and draft.locked_by != actor_id and current_user.get("role") != "Admin":
        raise HTTPException(status_code=403, detail=f"This draft is locked by {draft.locked_by}. Only that user or an Admin can release it.")
    locked_user = draft.locked_by
    draft.is_locked = False
    draft.locked_by = None
    draft.locked_at = None
    
    log_edit(db, "draft", str(draft_id), "unlock", actor_id, {"previously_locked_by": locked_user})
    
    db.commit()
    return {"message": "Draft file unlocked", "success": True}


# FR-054: Dispatch Draft
@router.post("/drafts/{draft_id}/dispatch")
def dispatch_draft(
    draft_id: int,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Dispatches a draft document, moving it to the final Outward Register.
    
    Implements:
    - FR-054: Moves draft to outward_register, renames file to next sequential number (e.g. 004.doc),
      moves file to Outward/{Year}/{FolderID}/, removes draft record from database.
    """
    draft = db.query(models.DraftFile).filter(models.DraftFile.draft_id == draft_id).first()
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    ensure_draft_not_pending_deletion(db, draft_id)

    if draft.is_locked:
         raise HTTPException(status_code=400, detail="Cannot dispatch a draft that is currently locked for editing.")

    actor_id = current_user.get("user_id")
    year = draft.year
    folder_id = draft.folder_id
    
    # FR-055/FR-054: Assign official outward number only at dispatch.
    outward_no = get_next_outward_no(folder_id, year, db)

    # File rename & move (FR-054)
    # From: Drafts/{Year}/{FolderID}/draft-...{ext}
    # To: Outward/{Year}/{FolderID}/{OutwardNo}.{ext}
    old_relative_path = draft.file_path
    
    ext = os.path.splitext(old_relative_path)[1]
    if not ext:
        ext = ".docx"
    new_filename = f"{outward_no}{ext}"
    is_compressed = (ext.lower() == ".zip")
    
    try:
        new_relative_path, full_new_path = filesystem_utils.move_draft_to_outward(
            get_iodms_root_path(), old_relative_path, year, folder_id, new_filename
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to move document on disk: {str(e)}")
        
    full_old_path = os.path.join(get_iodms_root_path(), old_relative_path)

    # If there are multiple files (from direct upload), move them all
    new_attachment_paths = [new_relative_path]
    if draft.attachment_paths:
        supporting_paths = [p for p in draft.attachment_paths if p and p != draft.file_path]
        for idx, p in enumerate(supporting_paths):
            p_ext = os.path.splitext(p)[1]
            p_filename = f"{outward_no}_attachment_{idx+1}{p_ext}"
            
            try:
                new_p, _ = filesystem_utils.move_draft_to_outward(
                    get_iodms_root_path(), p, year, folder_id, p_filename
                )
                new_attachment_paths.append(new_p)
            except Exception:
                pass
    else:
        new_attachment_paths = [new_relative_path]
    
    # Perform file system move and rename
    if not os.path.exists(full_old_path) and not os.path.exists(full_new_path):
        # If draft file is missing, create it directly in Outward folder
        create_draft_document(full_new_path, {
            "outward_no": outward_no,
            "year": year,
            "issuing_date": datetime.date.today().isoformat(),
            "folder_id": folder_id,
            "template_type": draft.template_type,
            "prepared_by": draft.prepared_by,
            "address_to": draft.address_to,
            "cc_to": draft.cc_to,
            "subject": draft.subject,
            "remarks": draft.remarks,
            "document_body": draft.document_body
        }, db)

    stamp_outward_reference(full_new_path, {
        "outward_no": outward_no,
        "year": year,
        "issuing_date": draft.issuing_date.isoformat(),
        "folder_id": folder_id,
        "template_type": draft.template_type,
        "prepared_by": draft.prepared_by,
        "address_to": draft.address_to,
        "cc_to": draft.cc_to,
        "subject": draft.subject,
        "remarks": draft.remarks
    }, db)

    # Insert record into Outward Register
    new_outward = models.OutwardRegister(
        outward_no=outward_no,
        folder_id=folder_id,
        year=year,
        issuing_date=datetime.date.today(),
        address_to=draft.address_to,
        cc_to=draft.cc_to,
        subject=draft.subject,
        remarks=draft.remarks,
        prepared_by=draft.prepared_by,
        actioned_by=actor_id,
        document_path=new_attachment_paths[0] if new_attachment_paths else new_relative_path,
        attachment_paths=new_attachment_paths,
        template_type=draft.template_type,
        linked_documents=draft.linked_documents,
        document_body=draft.document_body,
        is_compressed=is_compressed
    )
    
    db.add(new_outward)
    
    # Log dispatch action
    log_edit(db, "draft", str(draft_id), "dispatch", actor_id)
    log_edit(db, "outward", f"{folder_id}:{year}:{outward_no}", "create", actor_id)
    
    # Sync links
    source_id = f"outward:{folder_id}:{year}:{outward_no}"
    sync_bidirectional_links(db, source_id, [], draft.linked_documents)
    
    db.delete(draft)  # delete draft record (FR-054)
    db.commit()

    return {"message": "Document dispatched successfully", "outward_no": outward_no, "success": True}


# FR-056: Discard Draft
@router.delete("/drafts/{draft_id}")
def discard_draft(
    draft_id: int,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Requests draft discarding.
    
    Implements:
    - FR-056: Creates deletion request in pending_deletions; draft remains visible but greyed out.
    """
    draft = db.query(models.DraftFile).filter(models.DraftFile.draft_id == draft_id).first()
    if not draft:
         raise HTTPException(status_code=404, detail="Draft not found")
    actor_id = current_user.get("user_id")

    existing = db.query(models.PendingDeletion).filter(
        models.PendingDeletion.source_table == "draft_files",
        models.PendingDeletion.record_id == str(draft_id),
        models.PendingDeletion.status == "Pending"
    ).first()
    if existing:
        return {"message": "Draft discard request is already pending Admin approval.", "success": True}

    new_del = models.PendingDeletion(
        source_table="draft_files",
        record_id=str(draft_id),
        requested_by=actor_id,
        status="Pending"
    )
    db.add(new_del)
    log_edit(db, "draft", str(draft_id), "discard", actor_id)
    db.commit()
    return {"message": "Draft discard requested. Awaiting Admin approval.", "success": True}


# FR-090, FR-091, FR-092, FR-093: View Outward Register
@router.get("/register")
def get_outward_register(
    year: Optional[str] = None,
    page: int = 1,
    limit: int = 20,
    search_folder_id: Optional[str] = None,
    search_prepared_by: Optional[str] = None,
    search_address_to: Optional[str] = None,
    search_subject: Optional[str] = None,
    search_status: Optional[str] = None,
    search_date_from: Optional[str] = None,
    search_date_to: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """Retrieves paginated outward logs filtered by year and advanced search keywords.
    
    Implements:
    - FR-090: Displays outward entries.
    - FR-091: Address To and CC To printed as comma-separated values in single cell.
    - FR-092: Year filter.
    - FR-093: Search filters: Folder ID, Prepared By, Address To (text search), Subject.
    - FR-095: Marks row with Pending Deletion badge.
    """
    query = db.query(models.OutwardRegister).filter(models.OutwardRegister.status != "Permanently Deleted")

    if year and year != "All":
        try:
            y = int(year)
            query = query.filter(models.OutwardRegister.year == y)
        except ValueError:
            pass

    # Search filters
    if search_folder_id:
        query = query.filter(models.OutwardRegister.folder_id == search_folder_id)
    if search_prepared_by:
        query = query.filter(models.OutwardRegister.prepared_by == search_prepared_by)
    if search_subject:
        query = query.filter(models.OutwardRegister.subject.ilike(f"%{search_subject}%"))
    if search_address_to:
        matching_addrs = db.query(models.AddressBook.address_id).filter(
            models.AddressBook.name.ilike(f"%{search_address_to}%")
        ).all()
        matching_ids = [m[0] for m in matching_addrs]
        if matching_ids:
            query = query.filter(models.OutwardRegister.address_to.overlap(matching_ids))
        else:
            query = query.filter(models.OutwardRegister.address_to.overlap([-1]))
    if search_status:
        query = query.filter(models.OutwardRegister.status == search_status)
    if search_date_from:
        try:
            d_from = datetime.datetime.strptime(search_date_from, "%Y-%m-%d").date()
            query = query.filter(models.OutwardRegister.issuing_date >= d_from)
        except ValueError:
            pass
    if search_date_to:
        try:
            d_to = datetime.datetime.strptime(search_date_to, "%Y-%m-%d").date()
            query = query.filter(models.OutwardRegister.issuing_date <= d_to)
        except ValueError:
            pass
    
    total = query.count()
    offset = (page - 1) * limit
    results = query.order_by(models.OutwardRegister.outward_no.desc()).offset(offset).limit(limit).all()

    # Get pending deletions
    pending_deletes = db.query(models.PendingDeletion).filter(
        models.PendingDeletion.source_table == "outward_register",
        models.PendingDeletion.status == "Pending"
    ).all()
    pending_keys = {pd.record_id for pd in pending_deletes} # format "folder_id:year:outward_no"

    # FR-090: Batch fetch relationships to prevent N+1 queries
    folder_ids = list(set([r.folder_id for r in results]))
    folders = db.query(models.FolderType).filter(models.FolderType.folder_id.in_(folder_ids)).all()
    folder_map = {f.folder_id: f.folder_name for f in folders}

    address_ids = set()
    for r in results:
        address_ids.update(r.address_to or [])
        address_ids.update(r.cc_to or [])
    
    address_map = {}
    if address_ids:
        addresses = db.query(models.AddressBook).filter(models.AddressBook.address_id.in_(list(address_ids))).all()
        address_map = {a.address_id: a.name for a in addresses}

    output = []
    for r in results:
        key = f"{r.folder_id}:{r.year}:{r.outward_no}"
        folder_name = folder_map.get(r.folder_id, "")

        address_to_names = [address_map[a_id] for a_id in (r.address_to or []) if a_id in address_map]
        cc_to_names = [address_map[c_id] for c_id in (r.cc_to or []) if c_id in address_map]

        output.append({
            "outward_no": r.outward_no,
            "folder_id": r.folder_id,
            "folder_name": folder_name,
            "year": r.year,
            "issuing_date": r.issuing_date.isoformat(),
            "address_to_names": address_to_names,
            "cc_to_names": cc_to_names,
            "subject": r.subject,
            "remarks": r.remarks,
            "prepared_by": r.prepared_by,
            "document_path": r.document_path,
            "template_type": r.template_type,
            "is_pending_deletion": key in pending_keys
        })

    return {
        "total": total,
        "page": page,
        "limit": limit,
        "results": output
    }


# FR-094: View Document
@router.get("/view-document")
def view_document(path: str, db: Session = Depends(get_db)):
    """Serves the document file directly (PDF or DOC/DOCX) for in-browser viewer.
    
    Implements:
    - FR-094: In-Browser Document View
    """
    root_path = os.path.abspath(get_iodms_root_path())
    full_path = os.path.abspath(os.path.join(root_path, path.lstrip("/\\")))
    
    # Path traversal check
    if os.path.commonpath([root_path, full_path]) != root_path:
        raise HTTPException(status_code=403, detail="Forbidden: Path traversal detected")
        
    if not os.path.exists(full_path):
        raise HTTPException(status_code=404, detail="File not found on server disk")

    ext = os.path.splitext(full_path)[1].lower()
    media_types = {
        ".pdf": "application/pdf",
        ".doc": "application/msword",
        ".rtf": "application/rtf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    return FileResponse(full_path, media_type=media_types.get(ext, "application/octet-stream"), filename=os.path.basename(full_path))


# FR-170b: Upload Signed Copy / Additional Attachments to Dispatched Outward Record
@router.post("/{folder_id}/{year}/{outward_no}/attachments")
def upload_outward_attachments(
    folder_id: str,
    year: int,
    outward_no: str,
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Uploads signed copies or additional attachments to an already dispatched record."""
    record = db.query(models.OutwardRegister).filter(
        models.OutwardRegister.folder_id == folder_id,
        models.OutwardRegister.year == year,
        models.OutwardRegister.outward_no == outward_no
    ).first()
    
    if not record:
        raise HTTPException(status_code=404, detail="Outward record not found.")

    if not files or (len(files) == 1 and files[0].filename == ""):
        raise HTTPException(status_code=400, detail="At least one file is required.")

    actor_id = current_user.get("user_id")
    relative_folder, full_folder = filesystem_utils.ensure_folder_path(
        get_iodms_root_path(), "Outward", year, folder_id
    )
    
    existing_paths = record.attachment_paths or []
    new_paths = []

    for idx, file in enumerate(files, start=1):
        if not file.filename:
            continue
        ext = os.path.splitext(file.filename)[1] or ".bin"
        filename = f"{outward_no}_attachment_{len(existing_paths) + idx}{ext}"
        relative_path = os.path.join(relative_folder, filename).replace("\\", "/")
        full_path = os.path.join(full_folder, filename)

        try:
            with open(full_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            final_abs_path, was_compressed = filesystem_utils.compress_file_if_large(full_path)
            if was_compressed:
                final_filename = os.path.basename(final_abs_path)
                relative_path = os.path.join(relative_folder, final_filename).replace("\\", "/")
            new_paths.append(relative_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to save attachment: {str(e)}")

    record.attachment_paths = existing_paths + new_paths
    log_edit(db, "outward", f"{folder_id}:{year}:{outward_no}", "attach", actor_id, {"files": new_paths})
    db.commit()
    
    return {"message": "Files attached successfully", "paths": new_paths, "success": True}

# FR-058: Edit Audit Log Endpoint
@router.get("/edit-log/{record_type}/{record_id}")
def get_edit_log(record_type: str, record_id: str, db: Session = Depends(get_db)):
    """Gets the edit history for a specific record."""
    logs = db.query(models.EditLog).filter(
        models.EditLog.record_type == record_type,
        models.EditLog.record_id == record_id
    ).order_by(models.EditLog.edited_at.desc()).all()
    
    result = []
    for log in logs:
        editor = db.query(models.User).filter(models.User.user_id == log.edited_by).first()
        result.append({
            "action": log.action,
            "edited_by": editor.name if editor else log.edited_by,
            "edited_at": log.edited_at.isoformat(),
            "changes": log.changes
        })
    return result


# FR-095: Soft-delete outward register entry
@router.delete("/{folder_id}/{year}/{outward_no}")
def delete_outward_record(
    folder_id: str,
    year: int,
    outward_no: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """Submits outward record deletion request.
    
    Implements:
    - FR-095: Logs request in pending_deletions.
    """
    key = f"{folder_id}:{year}:{outward_no}"
    actor_id = current_user.get("user_id")
    
    existing = db.query(models.PendingDeletion).filter(
        models.PendingDeletion.source_table == "outward_register",
        models.PendingDeletion.record_id == key,
        models.PendingDeletion.status == "Pending"
    ).first()
    
    if existing:
        return {"message": "Deletion request already pending.", "success": True}
        
    new_del = models.PendingDeletion(
        source_table="outward_register",
        record_id=key,
        requested_by=actor_id,
        status="Pending"
    )
    db.add(new_del)
    log_edit(db, "outward", key, "delete_request", actor_id)
    db.commit()
    return {"message": "Deletion request submitted to Admin.", "success": True}
