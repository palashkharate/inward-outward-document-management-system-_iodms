from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "IODMS_Word_Launcher_Setup_Guide.docx"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)


def set_font(run, name="Calibri", size=11, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_style(style, size, color, before, after, bold=True):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_step(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_font(run, name="Consolas", size=9)
    return paragraph


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    set_style(doc.styles["Heading 1"], 16, BLUE, 18, 10)
    set_style(doc.styles["Heading 2"], 13, BLUE, 14, 7)
    set_style(doc.styles["Heading 3"], 12, DARK_BLUE, 10, 5)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("IODMS Word Launcher Setup Guide")
    set_font(title_run, size=22, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("Open and save IODMS LAN drafts directly in Microsoft Word")
    set_font(subtitle_run, size=11, color=RGBColor(89, 89, 89))

    add_heading(doc, "Purpose")
    doc.add_paragraph(
        "The IODMS Word Launcher lets the Drafts and Dispatch screen open the same document stored in the LAN shared IODMS folder. "
        "When the officer saves in Microsoft Word, the backend copy changes immediately because Word is editing the shared file itself."
    )

    add_heading(doc, "Before You Start")
    add_bullet(doc, "Microsoft Word must be installed on the officer PC.")
    add_bullet(doc, "The officer must be able to open and save files in the shared IODMS folder.")
    add_bullet(doc, "In IODMS Admin > System Settings, LAN Shared IODMS Path must match the client-visible UNC path, for example \\\\Server\\IODMS_DATA.")
    add_bullet(doc, "The backend IODMS Root Path and the LAN Shared IODMS Path must refer to the same underlying folder.")
    doc.add_paragraph(
        "Main computer note: on the PC that runs the backend, LAN Shared IODMS Path can be the same local folder path as IODMS Root Path. "
        "For example: C:\\Users\\Palash\\Desktop\\inword outword folder\\IODMS_DATA. Other PCs must use a UNC share path instead."
    )

    add_heading(doc, "Install on This PC")
    doc.add_paragraph("Run the installer once for the Windows user who will use IODMS on this PC.")
    add_step(doc, "Open Windows PowerShell. You do not need to run it as Administrator.")
    add_step(doc, "Go to the IODMS application folder that contains the scripts folder.")
    add_step(doc, "Run the following command:")
    add_code(doc, "powershell.exe -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\install_iodms_word_launcher.ps1")
    add_step(doc, "Confirm the message: IODMS Word Launcher installed for this Windows user.")
    doc.add_paragraph(
        "This installer creates a per-user Windows protocol registration only. It does not modify the IODMS database or move any documents."
    )

    add_heading(doc, "Configure the LAN Path")
    add_step(doc, "Sign in as an IODMS Admin.")
    add_step(doc, "Open Admin > System Settings.")
    add_step(doc, "Set LAN Shared IODMS Path to the UNC share officers use, such as \\\\Server\\IODMS_DATA.")
    add_step(doc, "Save the settings, then reload Drafts and Dispatch.")

    add_heading(doc, "Test Direct Word Editing")
    add_step(doc, "Open Drafts and Dispatch and expand any available draft.")
    add_step(doc, "Click Open in Word. The draft is locked before Word is launched.")
    add_step(doc, "If Chromium asks to open IODMS Word Launcher, allow it. This normally occurs only on the first launch.")
    add_step(doc, "Edit the document in Microsoft Word, save it, and close Word.")
    add_step(doc, "Return to IODMS and click Release Lock.")
    doc.add_paragraph(
        "To confirm that save-back worked, click View Document or open the same UNC path in File Explorer. The saved revision should be present without re-uploading."
    )

    add_heading(doc, "Install on Other PCs")
    doc.add_paragraph("Repeat the same installation command once for each Windows user profile that uses IODMS.")
    add_bullet(doc, "A domain administrator can deploy the installer as a user logon script.")
    add_bullet(doc, "Each PC still needs Microsoft Word and read/write access to the LAN share.")
    add_bullet(doc, "Each browser user may need to approve the IODMS Word Launcher prompt once.")

    add_heading(doc, "Troubleshooting")
    add_heading(doc, "Draft list says 'Failed to load drafts'", level=2)
    doc.add_paragraph("Restart the backend after applying the LAN editing update, then refresh the browser. This is not a Word Launcher installation problem.")
    add_heading(doc, "Word does not open", level=2)
    add_bullet(doc, "Run the installer again under the current Windows user, not only under an administrator account.")
    add_bullet(doc, "Confirm the PC can open the displayed UNC path in File Explorer.")
    add_bullet(doc, "Confirm Microsoft Word is the default application for DOC, DOCX, or RTF files.")
    add_bullet(doc, "Allow the browser's IODMS Word Launcher external-protocol prompt.")
    add_heading(doc, "Word opens but cannot save", level=2)
    add_bullet(doc, "Check that the Windows user has Modify permission on the shared IODMS folder.")
    add_bullet(doc, "Make sure another officer does not already hold the IODMS draft lock.")
    add_bullet(doc, "Do not use Download Draft for LAN editing; use Open in Word so Word opens the shared file.")

    add_heading(doc, "Remove the Launcher")
    doc.add_paragraph("To remove the launcher for the current Windows user, run:")
    add_code(doc, "Remove-Item -Path 'HKCU:\\Software\\Classes\\iodms-word' -Recurse -Force")
    doc.add_paragraph("Removing the launcher does not delete drafts or affect the IODMS backend.")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("IODMS Word Launcher Setup Guide")
    set_font(footer_run, size=8, color=RGBColor(89, 89, 89))

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
